import os
import time
import warnings
from langchain_core.tools import tool
from docx import Document
from langchain_ollama import ChatOllama
from config import MODEL_NAME
from ddgs import DDGS
import wikipedia

warnings.filterwarnings("ignore")

try: 
    wikipedia.set_lang("sv")
except: 
    pass

@tool
def create_word_document(filename: str, content: str):
    """
    Skapar ett äkta Microsoft Word-dokument (.docx).
    Används när användaren ber om att 'skriva ett word-dokument', 'skapa en rapport', etc.
    Input:
    - filename: Namnet på filen (t.ex. 'skolarbete.docx')
    - content: Texten som ska stå i dokumentet.
    """
    try:
        # Säkerställ filändelsen
        if not filename.endswith(".docx"):
            filename += ".docx"
            
        doc = Document()
        doc.add_heading('Dokument skapat av Jarvis', 0)
        doc.add_paragraph(content)
        
        doc.save(filename)
        
        # open document so user can see it
        os.system(f"start {filename}")
        
        return f"Word-dokumentet '{filename}' har skapats och öppnats."
    except Exception as e:
        return f"Fel vid skapande av Word-fil: {e}"


@tool
def create_notes_document(content: str, filename: str = "Anteckningar"):
    """
    Skapar ett nytt Word-dokument med anteckningar och öppnar det.
    Perfekt för när användaren säger 'öppna word och skriv ned' eller 'skapa anteckningar'.
    
    Input:
    - content: Texten/anteckningarna som ska skrivas
    - filename: Valfritt filnamn (standard: 'Anteckningar')
    """
    try:
        import time
        
        # Säkerställ filnamn
        if not filename.endswith(".docx"):
            filename += ".docx"
        
        print(f"DEBUG: Skapar anteckningsdokument: {filename}", flush=True)
        
        # Create Word document
        doc = Document()
        doc.add_heading('Anteckningar', 0)
        doc.add_paragraph(content)
        doc.add_paragraph("\n---")
        doc.add_paragraph(f"Skapad av Enigma • {time.strftime('%Y-%m-%d %H:%M')}")
        
        doc.save(filename)
        print(f"DEBUG: Dokument sparat: {filename}", flush=True)
        
        # Open in Word
        os.system(f"start {filename}")
        print(f"DEBUG: Öppnade dokument med Word", flush=True)
        
        return f"✓ Anteckningsdokument '{filename}' skapat och öppnat i Word."
    except Exception as e:
        print(f"DEBUG: Fel i create_notes_document: {e}", flush=True)
        return f"Fel: {str(e)}"


@tool
def create_research_document(topic: str, filename: str = "Forskning"):
    """
    Söker information om ett ämne på nätet, sammanfattar det och skriver till Word-dokument.
    Användare kan säga: "Öppna word och skriv fakta om [ämne]"
    
    Input:
    - topic: Ämnet att söka efter (t.ex. "vad Trump gjorde under sin period")
    - filename: Valfritt filnamn (standard: 'Forskning')
    """
    try:
        import os.path
        from pathlib import Path
        
        print(f"DEBUG: Startar forskning om: {topic}", flush=True)
        
        # Get possible save locations in order of preference
        possible_folders = [
            str(Path.home() / "Documents"),  # Documents folder
            str(Path.home() / "Desktop"),     # Desktop folder
            str(Path.home()),                 # Home folder
            os.path.dirname(os.path.abspath(__file__))  # Project folder (fallback)
        ]
        
        # Find first writable folder
        save_folder = None
        for folder in possible_folders:
            try:
                if os.path.exists(folder):
                    # Test if we can write
                    test_file = os.path.join(folder, ".enigma_test")
                    with open(test_file, 'w') as f:
                        f.write("test")
                    os.remove(test_file)
                    save_folder = folder
                    print(f"DEBUG: Använder mapp: {folder}", flush=True)
                    break
            except Exception as e:
                print(f"DEBUG: Mapp {folder} är inte skrivbar: {e}", flush=True)
                continue
        
        # If no folder found, use project directory
        if not save_folder:
            save_folder = os.path.dirname(os.path.abspath(__file__))
            print(f"DEBUG: Ingen skrivbar mapp hittad, använder: {save_folder}", flush=True)
        
        # Säkerställ filnamn (ta bort specialtecken)
        if not filename.endswith(".docx"):
            filename += ".docx"
        
        # Sanitize filename
        filename = filename.replace("?", "").replace(":", "").replace("|", "").replace("<", "").replace(">", "")
        
        # Full path
        full_path = os.path.join(save_folder, filename)
        print(f"DEBUG: Sparar till: {full_path}", flush=True)
        
        # Search for information using DDGS
        print(f"DEBUG: Söker information...", flush=True)
        search_results = ""
        
        try:
            print("DEBUG: Testar DuckDuckGo...", flush=True)
            ddgs = DDGS()
            results = ddgs.text(topic, max_results=5)
            
            if results:
                for r in results:
                    title = r.get('title', '')
                    body = r.get('body', '')
                    if title and body:
                        search_results += f"Källa: {title}\nInfo: {body}\n\n"
                
                if search_results:
                    print(f"DEBUG: DDG hittade {len(results)} resultat.", flush=True)
        except Exception as e:
            print(f"DEBUG: DuckDuckGo misslyckades ({e})", flush=True)
        
        # Fallback to Wikipedia
        if not search_results:
            try:
                print("DEBUG: Testar Wikipedia...", flush=True)
                wiki = wikipedia.summary(topic, sentences=15, auto_suggest=True)
                if wiki: 
                    print("DEBUG: Wikipedia hittade info.", flush=True)
                    search_results = f"Wikipedia: {wiki}"
            except Exception as e:
                print(f"DEBUG: Wikipedia hittade inget ({e}).", flush=True)
        
        if not search_results:
            return f"Kunde inte hitta information om '{topic}'. Försök med ett annat ämne."
        
        print(f"DEBUG: Hittade information, sammanfattar...", flush=True)
        
        # Summarize with AI
        llm = ChatOllama(model=MODEL_NAME, temperature=0.3)
        
        summary_prompt = f"""Du är en expert på att sammanfatta information. 
        
Din uppgift: Sammanfatta följande information om '{topic}' på ett tydligt och organiserat sätt.

REGLER:
1. Skriv på svenska
2. Organisera i avsnitt med rubriker
3. Gör det lättläst och strukturerat
4. Behåll all viktig information
5. Endast sammanfattningen, INGENTING ANNAT

INFORMATION ATT SAMMANFATTA:
{search_results}

SAMMANFATTNING:"""
        
        response = llm.invoke(summary_prompt)
        summary = response.content.strip()
        
        print(f"DEBUG: Sammanfattning klar, skapar dokument...", flush=True)
        
        # Create Word document
        doc = Document()
        doc.add_heading(f'{topic.capitalize()}', 0)
        doc.add_paragraph(f"Forskning utförda: {time.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph("\n" + summary)
        doc.add_paragraph("\n" + "="*50)
        doc.add_paragraph(f"Skapad av Enigma Research Assistant")
        
        doc.save(full_path)
        print(f"DEBUG: Dokument sparat: {full_path}", flush=True)
        
        # Open in Word with full path
        os.system(f'start "" "{full_path}"')
        print(f"DEBUG: Öppnade dokument med Word", flush=True)
        
        return f"✓ Forskning om '{topic}' skapat och öppnat: {filename}"
    except Exception as e:
        print(f"DEBUG: Fel i create_research_document: {e}", flush=True)
        import traceback
        traceback.print_exc()
        return f"Fel: {str(e)}"